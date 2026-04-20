"""
Resume Parser — supports PDF, DOCX, and image (OCR) formats.
Returns raw text and per-section extraction with confidence scores.
"""
import os
import re
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ── Section header patterns ────────────────────────────────────────────────
SECTION_PATTERNS = {
    "experience": re.compile(
        r"\b(experience|work\s+history|employment|professional\s+background|career)\b",
        re.IGNORECASE,
    ),
    "education": re.compile(
        r"\b(education|academic|qualifications?|degrees?|university|college)\b",
        re.IGNORECASE,
    ),
    "skills": re.compile(
        r"\b(skills?|technical\s+skills?|competenc(ies|y)|technologies|tools?)\b",
        re.IGNORECASE,
    ),
    "projects": re.compile(
        r"\b(projects?|portfolio|personal\s+projects?|key\s+projects?)\b",
        re.IGNORECASE,
    ),
    "certifications": re.compile(
        r"\b(certifications?|licen(ses?|ces?)|credentials?|courses?|training)\b",
        re.IGNORECASE,
    ),
    "summary": re.compile(
        r"\b(summary|objective|profile|about\s+me|overview)\b",
        re.IGNORECASE,
    ),
}


def parse_pdf(file_path: str) -> dict:
    """Extract text from PDF using pdfplumber with pypdf2 fallback."""
    text = ""
    confidence = 0.0
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages_text = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                pages_text.append(page_text)
            text = "\n".join(pages_text)
            confidence = 0.9 if len(text.strip()) > 100 else 0.5
    except Exception as e:
        logger.warning(f"pdfplumber failed for {file_path}: {e}. Trying pypdf2.")
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages_text = [page.extract_text() or "" for page in reader.pages]
                text = "\n".join(pages_text)
                confidence = 0.7 if len(text.strip()) > 100 else 0.3
        except Exception as e2:
            logger.error(f"PDF parsing completely failed for {file_path}: {e2}")
            return {"text": "", "confidence": 0.0, "error": str(e2)}

    return {"text": text, "confidence": confidence}


def parse_docx(file_path: str) -> dict:
    """Extract text from DOCX preserving paragraph order."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        text = "\n".join(paragraphs)
        confidence = 0.95 if len(text.strip()) > 100 else 0.5
        return {"text": text, "confidence": confidence}
    except Exception as e:
        logger.error(f"DOCX parsing failed for {file_path}: {e}")
        return {"text": "", "confidence": 0.0, "error": str(e)}


def parse_image(file_path: str) -> dict:
    """OCR-based text extraction from image files."""
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(file_path)
        # Preprocess: convert to grayscale for better OCR accuracy
        img = img.convert("L")
        text = pytesseract.image_to_string(img, config="--psm 6")
        confidence = 0.65 if len(text.strip()) > 50 else 0.3
        return {"text": text, "confidence": confidence}
    except Exception as e:
        logger.error(f"Image OCR failed for {file_path}: {e}")
        return {"text": "", "confidence": 0.0, "error": str(e)}


def parse_resume(file_path: str) -> dict:
    """
    Dispatch to the correct parser based on file extension.
    Returns: {text, confidence, sections, file_type}
    """
    ext = Path(file_path).suffix.lower()
    dispatch = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".jpg": parse_image,
        ".jpeg": parse_image,
        ".png": parse_image,
    }

    parser_fn = dispatch.get(ext)
    if not parser_fn:
        return {
            "text": "",
            "confidence": 0.0,
            "error": f"Unsupported file type: {ext}",
            "file_type": ext,
        }

    result = parser_fn(file_path)
    result["file_type"] = ext
    result["sections"] = extract_sections(result.get("text", ""))
    return result


def extract_sections(text: str) -> dict:
    """
    Split raw resume text into named sections.
    Returns dict mapping section_name -> text content.
    """
    if not text:
        return {}

    lines = text.split("\n")
    sections: dict = {}
    current_section = "header"
    buffer: list = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            buffer.append("")
            continue

        matched_section = None
        for section_name, pattern in SECTION_PATTERNS.items():
            # A section header is usually short (< 60 chars) and matches pattern
            if len(stripped) < 60 and pattern.search(stripped):
                matched_section = section_name
                break

        if matched_section:
            # Save previous buffer to current section
            if buffer:
                existing = sections.get(current_section, "")
                sections[current_section] = (existing + "\n" + "\n".join(buffer)).strip()
            current_section = matched_section
            buffer = []
        else:
            buffer.append(line)

    # Flush last buffer
    if buffer:
        existing = sections.get(current_section, "")
        sections[current_section] = (existing + "\n" + "\n".join(buffer)).strip()

    return sections
